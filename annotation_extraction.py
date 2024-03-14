import tkinter
from tkinter import ttk
from pathlib import Path
import fitz
import docx
from docx.shared import Pt
from tkinterdnd2 import DND_FILES, TkinterDnD

# I changed this value from 0.9 (original) to 0.1
_threshold_intersection = 0.1  # if the intersection is large enough.


def _check_contain(r_word: fitz.Rect, points: any) -> bool:
    """If `r_word` is contained in the rectangular area.

    The area of the intersection should be large enough compared to the
    area of the given word.

    Args:
        r_word (fitz.Rect): rectangular area of a single word.
        points (list): list of points in the rectangular area of the
            given part of a highlight.

    Returns:
        bool: whether `r_word` is contained in the rectangular area.
    """
    # `r` is mutable, so everytime a new `r` should be initiated.
    r = fitz.Quad(points).rect
    r.intersect(r_word)

    if r.getArea() >= r_word.getArea() * _threshold_intersection:
        contain = True
    else:
        contain = False
    return contain


def _extract_annot(annot: fitz.Annot, words_on_page: list) -> str:
    """Extract words in a given highlight.

    Args:
        annot (fitz.Annot): [description]
        words_on_page (list): [description]

    Returns:
        str: words in the entire highlight.
    """
    quad_points = annot.vertices
    quad_count = int(len(quad_points) / 4)
    sentences = ['' for i in range(quad_count)]
    for i in range(quad_count):
        points = quad_points[i * 4: i * 4 + 4]
        words = [
            w for w in words_on_page if
            _check_contain(fitz.Rect(w[:4]), points)
        ]
        sentences[i] = ' '.join(w[4] for w in words)
    sentence = ' '.join(sentences)

    return sentence


def get_title(file_path: Path) -> str:
    """
    Split the file path on `\\` in order to obtain the file title.
    :param file_path: a `Path` object from `pathlib` library.
    :return: a `string` containing the `title` of the file or a `string`
        indicating the `.pdf` extension was not found.
    """
    # splits file path into a list, then cycles through each item and finds .pdf extension
    # this can be a way to display title of pdf file given to application
    name = file_path.name           # gives file path to file and file title.pdf
    for item in name.split('\\'):
        find_extension = item.find('.pdf')  # finds part of a string in a string
        if find_extension == -1:    # if search criteria ('pdf' here) not found, returns -1
            return ".pdf extension not found"
        else:
            return item


def drop_inside_entry_box(event):
    """`Deletes` any previous entries in the `Entry` box and `inserts` the `event.data` (file path) with {} stripped."""
    entry_box.delete(0, tkinter.END)
    entry_box.insert("end", str(event.data).strip("{}"))


def call_back(event) -> None:
    """
    Pass a tkinter event as the argument. When the event is activated, the `save_file_entry` Entry box text is deleted.
    """
    save_file_entry.delete(0, "end")


def default_save_location():
    """
    Set the save location StringVar to a specific value.
    """



def extraction_to_doc() -> None:
    """
    `Extract` `PDF annotations` and save them to a `docx` Word document.
    """
    pdf_path = Path(entry_string.get().strip("{}"))

    try:
        title = get_title(pdf_path)
        pdf_doc = fitz.Document(pdf_path)              # opening pdf document
        word_doc = docx.Document()
        style = word_doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        word_doc.add_heading(title, 0)
        for pdf_page in pdf_doc:                        # cycle through each page in pdf
            page = pdf_page
            annot = page.first_annot                    # first annotation of page
            if annot is not None:
                page_text_page = page.get_textpage()        # allows text/images to be extracted
                page_text = page_text_page.extractWORDS()
                while annot:                                # while an annotation is found
                    if annot.type[0] in (8, 9, 10, 11):
                        annotation_type = f"Annotation Type: {annot.type[1]}"
                        page_number = f"Page: {str(page.number + 1)}"
                        word_doc.add_heading(annotation_type, 2)
                        word_doc.add_heading(page_number, 2)
                        extract_annotation = _extract_annot(annot, page_text)   # extract annotation
                        word_doc.add_paragraph(extract_annotation, style='Normal')
                    annot = annot.next
        pdf_doc.close()

        string_path = str(path)            # make file path a string
        split_path = string_path.split('\\')    # split string on \
        desktop = f"{split_path[0]}\\{split_path[1]}\\{split_path[2]}\\Desktop\\"   # get correct components for windows
        doc_save_title = save_entry_string.get()
        if doc_save_title in ('', "Save annotation file as..."):
            save_entry_string.set("Annotations")
            default_save_title = save_entry_string.get()
            word_doc.save(desktop + default_save_title + ".docx")
        else:
            word_doc.save(desktop + doc_save_title + ".docx")              # save doc to desktop
        end_label_string.set("Annotations have been successfully extracted")
    except (TypeError, RuntimeError):
        end_label_string.set("This is not a PDF document (ie. .pdf extension cannot be found in the file path).")


if __name__ == "__main__":

    # main window
    root_window = TkinterDnD.Tk()
    root_window.minsize(440, 115)
    root_window.maxsize(455, 125)
    background = '#e5e1e1'
    root_window.configure(background=background)
    root_window.title("PDF Annotation Extraction")

    # string variables
    entry_string = tkinter.StringVar()
    entry_string.set("Drag PDF here...")

    save_entry_string = tkinter.StringVar()
    save_entry_string.set("Save annotation file as...")

    label_string = tkinter.StringVar()
    label_string.set("Please drag a PDF file into the box below: ")

    end_label_string = tkinter.StringVar()
    end_label_string.set("Press the button to extract annotations")

    # col config
    root_window.columnconfigure(0, weight=1)
    root_window.columnconfigure(1, weight=1)

    # row config
    root_window.rowconfigure(0, weight=1)
    root_window.rowconfigure(1, weight=1)
    root_window.rowconfigure(2, weight=1)
    root_window.rowconfigure(3, weight=1)

    # widget style
    widget_style = ttk.Style(root_window)
    widget_style.configure('TButton', font=('Calibri', 12, 'bold'), background='blue')
    widget_style.configure('TLabel', font=('Calibri', 12, 'bold'))
    widget_style.configure('TEntry', font=('Calibri', 12, 'bold'), background='blue')

    # opening Label
    message_label = ttk.Label(root_window, textvariable=label_string, background=background)
    message_label.grid(row=0, column=0, sticky='w', padx=5)

    # error Label
    end_result_label = ttk.Label(root_window, textvariable=end_label_string, background=background, wraplength=300, justify=tkinter.LEFT)
    end_result_label.grid(row=3, column=0, sticky='w', padx=5)

    # Entry Box to drag and drop PDF
    entry_box = ttk.Entry(root_window, textvariable=entry_string)
    entry_box.grid(row=1, column=0, columnspan=2, sticky='ew', padx=5, pady=2)
    entry_box.drop_target_register(DND_FILES)
    entry_box.dnd_bind("<<Drop>>", drop_inside_entry_box)

    # Entry Box to ask user what save file should be called
    save_file_entry = ttk.Entry(root_window, textvariable=save_entry_string)
    save_file_entry.bind("<Button-1>", call_back)
    save_file_entry.grid(row=2, column=0, columnspan=2, sticky='ew', padx=5, pady=2)

    # Button
    extract_button = ttk.Button(root_window, text="Extract to Word", command=extraction_to_doc)
    extract_button.grid(row=3, column=1, sticky='e', padx=5, pady=2)

    tkinter.mainloop()
